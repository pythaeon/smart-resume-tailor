
from tkinter import Tk, Label, Button, filedialog, messagebox
import yaml
import re
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path

# This function takes a YAML path and generates a styled resume as a .docx file
def render_resume_from_yaml(yaml_path):
    # Load YAML
    with open(yaml_path, "r", encoding="utf-8") as file:
        resume_yaml = yaml.safe_load(file)

    # Load the style template
    template_path = yaml_path.parent / "Resume Style Template 1.0.docx"
    if not template_path.exists():
        messagebox.showerror("Template Missing", f"Missing template: {template_path.name}")
        return

    doc = Document(str(template_path))

    # Clear document body
# Clears all existing content from the Word template
    def clear_document_body(document):
        body = document._element.body
        for child in list(body):
            body.remove(child)

    clear_document_body(doc)

# Adds a paragraph with optional *italic* or **bold** inline formatting support
    def add_markdown_styled(text, style_name):
        p = doc.add_paragraph(style=style_name)
        tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for token in tokens:
            if token.startswith("**") and token.endswith("**"):
                run = p.add_run(token[2:-2])
                run.bold = True
            elif token.startswith("*") and token.endswith("*"):
                run = p.add_run(token[1:-1])
                run.italic = True
            else:
                run = p.add_run(token)
        return p

    # Header
    header = resume_yaml["resume"]["header"]
    add_markdown_styled(header["name"], "Main_Header_Name")
    for label, value in header["contact"].items():
        paragraph = doc.add_paragraph(style="Main_Header_Info")
        paragraph.add_run(f"{label.title()}: {value}")
    doc.add_paragraph()

    # Start With Why & Who I Am
    for key, title in [("start_with_why", "Start With Why"), ("who_i_am", "Who I Am")]:
        add_markdown_styled(title, "Section_Header")
        for para in resume_yaml["resume"][key]["content"].split("\n\n"):
            add_markdown_styled(para.strip(), "Section_Body")

    # How I Work
    add_markdown_styled("How I Work", "Section_Header")
    for block in resume_yaml["resume"]["how_i_work"]["blocks"]:
        if "block_heading" in block:
            add_markdown_styled(block["block_heading"], "Section_Subheader")
            add_markdown_styled(", ".join(block.get("list", [])), "Section_Body")
        elif "outro" in block:
            add_markdown_styled(block["outro"], "Section_Body")

    # Education
    add_markdown_styled("Education & Certification", "Section_Header")
    edu = resume_yaml["resume"]["education"]
    college = edu.get("college", {})
    cert = edu.get("certification", {})
    if college:
        add_markdown_styled(
            f"{college.get('college_name')} — {college.get('degree_type')} in {college.get('major_name')} ({college.get('year')})",
            "Section_Body"
        )
    if cert:
        add_markdown_styled(
            f"{cert.get('certification type')} — {cert.get('certifying_body')} ({cert.get('year')})",
            "Section_Body"
        )

    # Experience
    doc.add_page_break()
    add_markdown_styled(resume_yaml["resume"]["experience"]["section_title"], "Section_Header")
    job_counter = 0
    for job in resume_yaml["resume"]["experience"]["jobs"]:
        add_markdown_styled(f"{job['job_title']} — {job['company_name']}", "Role_Heading")
        add_markdown_styled(f"{job['location']} | {job['date_start']} – {job['date_end']}", "Role_Loc_Dates")
        for para in job["summary"].split("\n\n"):
            add_markdown_styled(para.strip(), "Role_Summary")
        for bullet in job["bullets"]:
            add_markdown_styled(bullet, "Role_Bullets")
        if job["punch_quote"]:
            add_markdown_styled(job["punch_quote"], "Role_Punch_Quote")
        job_counter += 1
        if job_counter % 2 == 0 and job_counter != len(resume_yaml["resume"]["experience"]["jobs"]):
            doc.add_page_break()

    # Save .docx
    out_path = yaml_path.with_suffix(".docx")
    doc.save(out_path)
    messagebox.showinfo("Success", f"Resume saved as:\n{out_path.name}")

# GUI
# Launches file picker dialog and triggers resume rendering
def choose_file():
    file_path = filedialog.askopenfilename(filetypes=[("YAML Files", "*.yml *.yaml")])
    if file_path:
        render_resume_from_yaml(Path(file_path))

# Build GUI
# Build and launch the main GUI window
root = Tk()
root.title("Resume Renderer")
root.geometry("400x150")
Label(root, text="Resume Renderer", font=("Arial", 14)).pack(pady=10)
Button(root, text="Choose Resume YAML", command=choose_file, width=30).pack(pady=5)
Label(root, text="(Make sure template is in the same folder)", font=("Arial", 9)).pack()
root.mainloop()
